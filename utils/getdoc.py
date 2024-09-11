import sys
import os

current_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.abspath(os.path.join(current_dir, os.pardir))
sys.path.insert(0, parent_dir)


import validator_functions
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_bolded_text(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True

def add_function_doc(doc, func):
    doc.add_heading(func.__name__, level=2)
    if func.__doc__:
        for line in func.__doc__.splitlines():
            if 'Parameters' in line:
                paragraph = doc.add_paragraph()
                add_bolded_text(paragraph, line.strip())
            elif '-' in line:
                param, desc = line.split('-', 1)
                paragraph = doc.add_paragraph()
                add_bolded_text(paragraph, param.strip())
                paragraph.add_run(' - ' + desc.strip())
            else:
                doc.add_paragraph(line.strip())

def create_document(module):
    doc = Document()
    doc.add_heading('Module Documentation', 0)

    for name, func in module.__dict__.items():
        if callable(func):
            add_function_doc(doc, func)

    doc.save('documentation.docx')

create_document(validator_functions)
